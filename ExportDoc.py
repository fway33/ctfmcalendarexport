from docx import Document

from DataStructs import locations, location_exceptions, degrees, full_list_calendar_entries, dinner


def create_word_doc():
    print("=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=-=")
    print(full_list_calendar_entries)
    print(locations)

    document = Document()
    document.add_heading("Calendar Events")

    array_idx = 0;
    for item in full_list_calendar_entries:
        print(item)
        if item[0] == -1:
            print(array_idx)
            print(": ")
            print(item)
            continue

        p = document.add_paragraph()

        # for this paragraph (an event) first we want to see if there is a
        # dinner and/or degree.  Check those arrays first to see if the array_idx matches
        if array_idx in degrees:
            p.add_run('*** DEGREE *** ')
        if array_idx in dinner:
            p.add_run(" DINNER ")

        # Now put the time in, in bold:
        p.add_run(item[4]).bold = True
        p.add_run(', ')

        # Need the lodge name
        p.add_run(item[0])
        run = p.add_run()
        run.add_break()

        print('..........')
        print(array_idx)

        # Next up is location.
        if locations[array_idx]:
            print(locations[array_idx])
            p.add_run(locations[array_idx])
        print(item[3])
        p.add_run('| ')
        p.add_run(item[3])
        p.add_run(' |')

        # Next up event title
        p.add_run(item[1])
        run.add_break
        # event descr
        p.add_run(item[2])
        

        print('@@@@@@@@@@@@@@@@@')
        array_idx += 1

    document.save('scratch.docx')
