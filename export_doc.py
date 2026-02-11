import logging
from docx import Document

from data_structs import Col, lodge_locations, location_exceptions, degrees, full_list_calendar_entries, has_dinner

logger = logging.getLogger(__name__)

def create_word_doc() -> None:
    """ This will create the base docx file that will be tweaked and then put into the paper"""
    #create a document object
    document = Document()
    document.add_heading("Calendar Events")

    # full_list_calendar_entries, a list of lists,  holds all the massaged data from the cells.
    # We will loop through this and write out lines in to the Word doc.  The idea is
    # that most of them will be very close to the final format required.  It will be
    # easy to tweak locations etc.
    array_idx = 0

    [add_event_to_doc(document,array_idx, event_row) for event_row in full_list_calendar_entries if event_row[0] != -1]

    # After writing all that, save the document.
    document.save('scratch.docx')
    logger.debug(str(location_exceptions))


def add_event_to_doc(document: Document, array_idx: int, event_row: list[any]) -> None:
        # Each paragraph is one event.
        p = document.add_paragraph()

        # For this paragraph (an event) first we want to see if there is a
        # dinner and/or degree.  Check those arrays first to see if the current array_idx matches
        # anything from the dinner or degrees lists.
        if array_idx in degrees:
            p.add_run('*** DEGREE *** ')
        if array_idx in has_dinner:
            p.add_run(" DINNER ")

        # Now put the time in, in bold:
        p.add_run(event_row[Col.DATE.value]).bold = True
        p.add_run(', ')

        # Need the lodge name
        p.add_run(event_row[Col.LODGE.value])
        run = p.add_run()
        run.add_break()

        # Next up is location.  If there is something in the locations list for this array index
        # print that.  Then put in the full location from the full_list_calendar_entries, delimited by '|'
        if lodge_locations[array_idx]:
            p.add_run(lodge_locations[array_idx])
        p.add_run('| ')
        p.add_run(event_row[Col.LOCATION.value])
        p.add_run(' |')

        # Next up event title
        p.add_run(event_row[Col.TITLE.value])
        run.add_break()
        # And finally the event description.
        p.add_run(event_row[Col.DESCR.value])

        array_idx += 1

